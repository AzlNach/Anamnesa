#!/usr/bin/env python3
"""
Google Drive Indexer for RAG System
Mengekstrak, chunk, dan embed dokumen dari Google Drive untuk sistem RAG Anamnesa
"""

import os
import json
import logging
from pathlib import Path
from typing import List, Dict, Any
import tempfile
import mimetypes

# Google APIs
from google.auth.transport.requests import Request
from google.oauth2.credentials import Credentials
from google_auth_oauthlib.flow import InstalledAppFlow
from googleapiclient.discovery import build
from googleapiclient.errors import HttpError
import io
from googleapiclient.http import MediaIoBaseDownload

# Document processing
import PyPDF2
from docx import Document
import pandas as pd
from PIL import Image

# LangChain and AI
from langchain.text_splitter import RecursiveCharacterTextSplitter
import google.generativeai as genai
import numpy as np
from sklearn.metrics.pairwise import cosine_similarity

# Configuration
SCOPES = ['https://www.googleapis.com/auth/drive.readonly']
CREDENTIALS_FILE = 'credentials.json'  # Changed from './rag-system/credentials.json'
TOKEN_FILE = 'token.json'  # Changed from './rag-system/token.json'

# Setup logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class GoogleDriveIndexer:
    """Kelas untuk indexing dokumen dari Google Drive"""
    
    def __init__(self, gemini_api_key: str, credentials_file: str = CREDENTIALS_FILE):
        self.credentials_file = credentials_file
        self.token_file = TOKEN_FILE
        self.service = None
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            length_function=len,
        )
        
        # Setup Gemini API
        genai.configure(api_key=gemini_api_key)
        self.embedding_model = "models/text-embedding-004"
        
        # File processing cache to avoid duplicates
        self.processed_files = set()
        self.cache_file = Path(__file__).parent / 'data' / 'indexer_cache.json'
        
        # Load existing cache
        self._load_cache()
        
        # Initialize Google Drive service
        self._authenticate()
    
    def _load_cache(self):
        """Load cache of previously processed files"""
        try:
            if self.cache_file.exists():
                with open(self.cache_file, 'r', encoding='utf-8') as f:
                    cache_data = json.load(f)
                    self.processed_files = set(cache_data.get('processed_files', []))
                logger.info(f"Loaded cache: {len(self.processed_files)} previously processed files")
            else:
                logger.info("No cache file found, starting fresh")
        except Exception as e:
            logger.warning(f"Error loading cache: {e}")
            self.processed_files = set()
    
    def _save_cache(self):
        """Save cache of processed files"""
        try:
            self.cache_file.parent.mkdir(parents=True, exist_ok=True)
            cache_data = {
                'processed_files': list(self.processed_files),
                'last_updated': str(Path(__file__).parent)
            }
            with open(self.cache_file, 'w', encoding='utf-8') as f:
                json.dump(cache_data, f, ensure_ascii=False, indent=2)
            logger.info(f"Cache saved: {len(self.processed_files)} processed files")
        except Exception as e:
            logger.error(f"Error saving cache: {e}")
    
    def _get_file_hash(self, file_data: Dict) -> str:
        """Generate unique hash for file based on ID and modification time"""
        # Use file ID and size as unique identifier
        file_id = file_data.get('id', '')
        file_size = file_data.get('size', 0)
        file_name = file_data.get('name', '')
        return f"{file_id}_{file_size}_{file_name}"
    
    def _authenticate(self):
        """Authenticate with Google Drive API"""
        creds = None
        
        # Load existing token
        if os.path.exists(self.token_file):
            creds = Credentials.from_authorized_user_file(self.token_file, SCOPES)
        
        # If there are no (valid) credentials available, let the user log in
        if not creds or not creds.valid:
            if creds and creds.expired and creds.refresh_token:
                creds.refresh(Request())
            else:
                flow = InstalledAppFlow.from_client_secrets_file(
                    self.credentials_file, SCOPES)
                creds = flow.run_local_server(port=0)
            
            # Save the credentials for the next run
            with open(self.token_file, 'w') as token:
                token.write(creds.to_json())
        
        try:
            self.service = build('drive', 'v3', credentials=creds)
            logger.info("Successfully authenticated with Google Drive API")
        except HttpError as error:
            logger.error(f"An error occurred: {error}")
            self.service = None
    
    def list_files(self, folder_id: str = None, file_types: List[str] = None, recursive: bool = True) -> List[Dict]:
        """List files from Google Drive"""
        if not self.service:
            logger.error("Google Drive service not initialized")
            return []
        
        try:
            all_files = []
            
            # Build query
            query_parts = []
            
            if folder_id:
                if recursive:
                    # Get all files in folder and subfolders
                    all_files.extend(self._get_files_recursive(folder_id, file_types))
                    return all_files
                else:
                    query_parts.append(f"'{folder_id}' in parents")
            
            if file_types:
                mime_queries = []
                for file_type in file_types:
                    if file_type == 'pdf':
                        mime_queries.append("mimeType='application/pdf'")
                    elif file_type == 'docx':
                        mime_queries.append("mimeType='application/vnd.openxmlformats-officedocument.wordprocessingml.document'")
                    elif file_type == 'txt':
                        mime_queries.append("mimeType='text/plain'")
                    elif file_type == 'csv':
                        mime_queries.append("mimeType='text/csv'")
                
                if mime_queries:
                    query_parts.append(f"({' or '.join(mime_queries)})")
            
            query = ' and '.join(query_parts) if query_parts else None
            
            results = self.service.files().list(
                q=query,
                fields="nextPageToken, files(id, name, mimeType, size)"
            ).execute()
            
            files = results.get('files', [])
            logger.info(f"Found {len(files)} files in Google Drive")
            return files
            
        except HttpError as error:
            logger.error(f"An error occurred while listing files: {error}")
            return []
    
    def _get_files_recursive(self, folder_id: str, file_types: List[str] = None) -> List[Dict]:
        """Recursively get all files from folder and subfolders"""
        all_files = []
        
        try:
            # Get files in current folder
            query_parts = [f"'{folder_id}' in parents"]
            
            if file_types:
                mime_queries = []
                for file_type in file_types:
                    if file_type == 'pdf':
                        mime_queries.append("mimeType='application/pdf'")
                    elif file_type == 'docx':
                        mime_queries.append("mimeType='application/vnd.openxmlformats-officedocument.wordprocessingml.document'")
                    elif file_type == 'txt':
                        mime_queries.append("mimeType='text/plain'")
                    elif file_type == 'csv':
                        mime_queries.append("mimeType='text/csv'")
                
                if mime_queries:
                    query_parts.append(f"({' or '.join(mime_queries)})")
            
            # Get files
            file_query = ' and '.join(query_parts)
            results = self.service.files().list(
                q=file_query,
                fields="files(id, name, mimeType, size)"
            ).execute()
            
            files = results.get('files', [])
            all_files.extend(files)
            
            # Get subfolders
            folder_query = f"'{folder_id}' in parents and mimeType='application/vnd.google-apps.folder'"
            folder_results = self.service.files().list(
                q=folder_query,
                fields="files(id, name)"
            ).execute()
            
            subfolders = folder_results.get('files', [])
            
            # Recursively process subfolders
            for subfolder in subfolders:
                logger.info(f"Processing subfolder: {subfolder['name']}")
                subfolder_files = self._get_files_recursive(subfolder['id'], file_types)
                all_files.extend(subfolder_files)
            
            return all_files
            
        except HttpError as error:
            logger.error(f"Error in recursive file search: {error}")
            return []
    
    def download_file(self, file_id: str, file_name: str) -> str:
        """Download file from Google Drive and return content"""
        if not self.service:
            logger.error("Google Drive service not initialized")
            return ""
        
        try:
            request = self.service.files().get_media(fileId=file_id)
            file_content = io.BytesIO()
            downloader = MediaIoBaseDownload(file_content, request)
            
            done = False
            while done is False:
                status, done = downloader.next_chunk()
            
            file_content.seek(0)
            return file_content.getvalue()
            
        except HttpError as error:
            logger.error(f"An error occurred while downloading file {file_name}: {error}")
            return None
    
    def extract_text_from_pdf(self, content: bytes, file_name: str = "unknown") -> str:
        """Extract text from PDF content with progress tracking"""
        try:
            pdf_file = io.BytesIO(content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            total_pages = len(pdf_reader.pages)
            
            logger.info(f"[{file_name}] Starting PDF extraction: {total_pages} pages")
            
            text = ""
            for page_num in range(total_pages):
                page = pdf_reader.pages[page_num]
                page_text = page.extract_text()
                text += page_text + "\n"
                
                # Progress indicator every 10 pages or at the end
                if (page_num + 1) % 10 == 0 or page_num == total_pages - 1:
                    progress_pct = ((page_num + 1) / total_pages) * 100
                    chars_extracted = len(text)
                    logger.info(f"[{file_name}] PDF progress: {page_num + 1}/{total_pages} pages ({progress_pct:.1f}%) - {chars_extracted:,} characters")
            
            logger.info(f"[{file_name}] PDF extraction completed: {len(text):,} total characters")
            return text
        except Exception as e:
            logger.error(f"[{file_name}] Error extracting text from PDF: {e}")
            return ""
    
    def extract_text_from_docx(self, content: bytes, file_name: str = "unknown") -> str:
        """Extract text from DOCX content with progress tracking"""
        try:
            docx_file = io.BytesIO(content)
            doc = Document(docx_file)
            total_paragraphs = len(doc.paragraphs)
            
            logger.info(f"[{file_name}] Starting DOCX extraction: {total_paragraphs} paragraphs")
            
            text = ""
            for i, paragraph in enumerate(doc.paragraphs):
                text += paragraph.text + "\n"
                
                # Progress indicator every 100 paragraphs or at the end
                if (i + 1) % 100 == 0 or i == total_paragraphs - 1:
                    progress_pct = ((i + 1) / total_paragraphs) * 100
                    chars_extracted = len(text)
                    logger.info(f"[{file_name}] DOCX progress: {i + 1}/{total_paragraphs} paragraphs ({progress_pct:.1f}%) - {chars_extracted:,} characters")
            
            logger.info(f"[{file_name}] DOCX extraction completed: {len(text):,} total characters")
            return text
        except Exception as e:
            logger.error(f"[{file_name}] Error extracting text from DOCX: {e}")
            return ""
    
    def extract_text_from_txt(self, content: bytes) -> str:
        """Extract text from TXT content"""
        try:
            return content.decode('utf-8')
        except Exception as e:
            logger.error(f"Error extracting text from TXT: {e}")
            return ""
    
    def extract_text_from_csv(self, content: bytes, file_name: str = "unknown") -> str:
        """Extract text from CSV content without row limit and with progress tracking"""
        try:
            csv_file = io.BytesIO(content)
            df = pd.read_csv(csv_file)
            
            total_rows = len(df)
            logger.info(f"[{file_name}] Starting CSV extraction: {total_rows:,} rows, {len(df.columns)} columns")
            
            # Convert dataframe to text format
            text = f"CSV Data ({total_rows:,} rows):\n"
            text += f"Columns: {', '.join(df.columns.tolist())}\n\n"
            
            # Add each row as text with progress indication
            for index, row in df.iterrows():
                if (index + 1) % 1000 == 0:  # Progress indicator every 1000 rows
                    progress_pct = ((index + 1) / total_rows) * 100
                    chars_extracted = len(text)
                    logger.info(f"[{file_name}] CSV progress: {index + 1:,}/{total_rows:,} rows ({progress_pct:.1f}%) - {chars_extracted:,} characters")
                
                row_text = ", ".join([f"{col}: {val}" for col, val in row.items() if pd.notna(val)])
                text += f"Row {index + 1}: {row_text}\n"
            
            logger.info(f"[{file_name}] CSV extraction completed: {total_rows:,} rows, {len(text):,} characters")
            return text
        except Exception as e:
            logger.error(f"[{file_name}] Error extracting text from CSV: {e}")
            return ""
    
    def extract_text_from_json(self, content: bytes) -> str:
        """Extract text from JSON content"""
        try:
            json_text = content.decode('utf-8')
            data = json.loads(json_text)
            
            def json_to_text(obj, prefix=""):
                """Recursively convert JSON object to readable text"""
                text = ""
                if isinstance(obj, dict):
                    for key, value in obj.items():
                        if isinstance(value, (dict, list)):
                            text += f"{prefix}{key}:\n"
                            text += json_to_text(value, prefix + "  ")
                        else:
                            text += f"{prefix}{key}: {value}\n"
                elif isinstance(obj, list):
                    for i, item in enumerate(obj):
                        text += f"{prefix}Item {i + 1}:\n"
                        text += json_to_text(item, prefix + "  ")
                else:
                    text += f"{prefix}{obj}\n"
                return text
            
            text = "JSON Data:\n"
            text += json_to_text(data)
            return text
            
        except Exception as e:
            logger.error(f"Error extracting text from JSON: {e}")
            return ""
    
    def extract_text_from_file(self, file_data: Dict, content: bytes) -> str:
        """Extract text from file based on mime type with progress tracking"""
        mime_type = file_data.get('mimeType', '')
        file_name = file_data.get('name', 'unknown')
        
        logger.info(f"[{file_name}] Starting text extraction ({mime_type})")
        
        if mime_type == 'application/pdf':
            return self.extract_text_from_pdf(content, file_name)
        elif mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
            return self.extract_text_from_docx(content, file_name)
        elif mime_type == 'text/plain':
            text = self.extract_text_from_txt(content)
            logger.info(f"[{file_name}] TXT extraction completed: {len(text):,} characters")
            return text
        elif mime_type == 'text/csv':
            return self.extract_text_from_csv(content, file_name)
        elif mime_type == 'application/json' or file_name.lower().endswith('.json'):
            text = self.extract_text_from_json(content)
            logger.info(f"[{file_name}] JSON extraction completed: {len(text):,} characters")
            return text
        else:
            logger.warning(f"[{file_name}] Unsupported file type: {mime_type}")
            return ""
    
    def create_embeddings(self, texts: List[str], file_name: str = "unknown") -> List[List[float]]:
        """Create embeddings using Gemini API with progress tracking"""
        embeddings = []
        total_texts = len(texts)
        
        logger.info(f"🧠 [{file_name}] Creating embeddings for {total_texts} text chunks...")
        
        for i, text in enumerate(texts):
            try:
                result = genai.embed_content(
                    model=self.embedding_model,
                    content=text,
                    task_type="retrieval_document"
                )
                embeddings.append(result['embedding'])
                
                # Progress indicator every 10 embeddings or at the end
                if (i + 1) % 10 == 0 or i == total_texts - 1:
                    progress_pct = ((i + 1) / total_texts) * 100
                    logger.info(f"🧠 [{file_name}] Embedding progress: {i + 1}/{total_texts} ({progress_pct:.1f}%)")
                    
            except Exception as e:
                logger.error(f"🧠 [{file_name}] Error creating embedding for chunk {i + 1}: {e}")
                # Create zero embedding as fallback
                embeddings.append([0.0] * 768)  # Assuming 768 dimensions
        
        logger.info(f"✅ [{file_name}] Embeddings completed: {len(embeddings)} embeddings created")
        return embeddings
    
    def process_files(self, folder_id: str = None) -> List[Dict]:
        """Process all files and create embeddings with detailed progress tracking and deduplication"""
        # Supported file types - now includes JSON
        supported_types = ['pdf', 'docx', 'txt', 'csv', 'json']
        
        # Get list of files
        files = self.list_files(folder_id, supported_types)
        total_files = len(files)
        
        logger.info(f"🚀 Starting file processing: {total_files} files found")
        
        documents = []
        processed_files = 0
        skipped_files = 0
        total_chunks = 0
        total_characters = 0
        
        for file_index, file_data in enumerate(files, 1):
            file_name = file_data['name']
            file_size = file_data.get('size', 0)
            file_hash = self._get_file_hash(file_data)
            
            # Check if file already processed
            if file_hash in self.processed_files:
                logger.info(f"⏭️  Skipping already processed file {file_index}/{total_files}: {file_name}")
                skipped_files += 1
                continue
            
            logger.info(f"📄 Processing file {file_index}/{total_files}: {file_name} ({file_size} bytes)")
            
            # Download file content
            content = self.download_file(file_data['id'], file_data['name'])
            if not content:
                logger.warning(f"❌ Failed to download: {file_name}")
                continue
            
            # Extract text
            text = self.extract_text_from_file(file_data, content)
            if not text.strip():
                logger.warning(f"⚠️  No text extracted from file: {file_name}")
                # Still mark as processed to avoid retrying
                self.processed_files.add(file_hash)
                continue
            
            file_characters = len(text)
            total_characters += file_characters
            
            logger.info(f"📝 [{file_name}] Text extraction summary: {file_characters} characters")
            
            # Split text into chunks
            logger.info(f"🔧 [{file_name}] Splitting text into chunks...")
            chunks = self.text_splitter.split_text(text)
            file_chunks = len(chunks)
            total_chunks += file_chunks
            
            logger.info(f"✂️  [{file_name}] Created {file_chunks} chunks")
            
            # Create embeddings for chunks
            logger.info(f"🧠 [{file_name}] Creating embeddings for {file_chunks} chunks...")
            embeddings = self.create_embeddings(chunks, file_name)
            
            # Store document data
            for i, (chunk, embedding) in enumerate(zip(chunks, embeddings)):
                doc = {
                    'id': f"{file_data['id']}_{i}",
                    'file_id': file_data['id'],
                    'file_name': file_data['name'],
                    'file_type': file_data['mimeType'],
                    'chunk_index': i,
                    'content': chunk,
                    'embedding': embedding,
                    'metadata': {
                        'file_size': file_data.get('size', 0),
                        'chunk_length': len(chunk),
                        'parent_folders': file_data.get('parents', []),
                        'file_hash': file_hash
                    }
                }
                documents.append(doc)
            
            # Mark file as processed
            self.processed_files.add(file_hash)
            processed_files += 1
            
            # Progress summary
            progress_pct = (file_index / total_files) * 100
            logger.info(f"✅ [{file_name}] Completed - {file_chunks} chunks, {file_characters} characters")
            logger.info(f"📊 Overall progress: {file_index}/{total_files} files ({progress_pct:.1f}%) - {total_chunks} total chunks, {total_characters} total characters")
            logger.info("-" * 80)
        
        # Save cache
        self._save_cache()
        
        # Final summary
        logger.info(f"🎉 Processing completed!")
        logger.info(f"📈 Final statistics:")
        logger.info(f"   Files found: {total_files}")
        logger.info(f"   Files processed: {processed_files}")
        logger.info(f"   Files skipped (already processed): {skipped_files}")
        logger.info(f"   Total chunks created: {total_chunks}")
        logger.info(f"   Total characters extracted: {total_characters}")
        if processed_files > 0:
            logger.info(f"   Average chunks per file: {total_chunks/processed_files:.1f}")
            logger.info(f"   Average characters per file: {total_characters/processed_files:.0f}")
        
        return documents
    
    def clear_cache(self):
        """Clear the processing cache"""
        self.processed_files.clear()
        if self.cache_file.exists():
            self.cache_file.unlink()
        logger.info("Cache cleared successfully")
    
    def get_cache_info(self):
        """Get information about cached files"""
        return {
            'cached_files_count': len(self.processed_files),
            'cache_file_exists': self.cache_file.exists(),
            'cache_file_path': str(self.cache_file)
        }
    
    def save_to_json(self, documents: List[Dict], output_file: str):
        """Save documents to JSON file"""
        try:
            # Create output directory if it doesn't exist
            output_path = Path(output_file)
            output_path.parent.mkdir(parents=True, exist_ok=True)
            
            with open(output_file, 'w', encoding='utf-8') as f:
                json.dump(documents, f, ensure_ascii=False, indent=2)
            
            logger.info(f"Saved {len(documents)} documents to {output_file}")
        except Exception as e:
            logger.error(f"Error saving documents to JSON: {e}")

def main():
    """Main function to run the indexer with testing capabilities"""
    print("🚀 Google Drive Indexer for RAG System")
    print("=" * 60)
    
    # Load environment variables from .env.local
    current_dir = Path(__file__).parent
    env_file = current_dir.parent / '.env.local'
    
    gemini_api_key = None
    if env_file.exists():
        with open(env_file, 'r') as f:
            for line in f:
                if line.startswith('GOOGLE_API_KEY='):
                    gemini_api_key = line.split('=', 1)[1].strip()
                    break
    
    if not gemini_api_key:
        # Fallback to dotenv if .env.local not found
        from dotenv import load_dotenv
        load_dotenv()
        gemini_api_key = os.getenv('GOOGLE_API_KEY')
    
    if not gemini_api_key:
        logger.error("❌ GOOGLE_API_KEY not found in .env.local or environment variables")
        return
    
    try:
        # Initialize indexer
        logger.info("🔧 Initializing Google Drive Indexer...")
        indexer = GoogleDriveIndexer(gemini_api_key)
        
        if not indexer.service:
            logger.error("❌ Failed to authenticate with Google Drive")
            return
        
        # Show cache info
        cache_info = indexer.get_cache_info()
        logger.info(f"📁 Cache info: {cache_info['cached_files_count']} files previously processed")
        
        # Interactive menu
        while True:
            print("\n" + "=" * 60)
            print("📋 Choose an option:")
            print("1. 🔍 Test indexer (process limited files)")
            print("2. 🚀 Full processing (process all files)")
            print("3. 📊 Show cache information")
            print("4. 🗑️  Clear cache")
            print("5. ❌ Exit")
            
            choice = input("\nEnter your choice (1-5): ").strip()
            
            if choice == '1':
                # Test mode - limit to first 3 files
                print("\n🧪 Testing mode - processing first 3 files")
                folder_id = input("📁 Enter Google Drive folder ID (or press Enter for all files): ").strip()
                folder_id = folder_id if folder_id else None
                
                # Get files and limit to 3 for testing
                files = indexer.list_files(folder_id, ['pdf', 'docx', 'txt', 'csv', 'json'])
                if len(files) > 3:
                    files = files[:3]
                    logger.info(f"🎯 Limited to first 3 files for testing")
                
                # Process files
                test_documents = []
                for file_data in files:
                    file_hash = indexer._get_file_hash(file_data)
                    if file_hash not in indexer.processed_files:
                        # Temporarily process single file
                        single_file_docs = indexer.process_files(None)  # Will use cached list
                        test_documents.extend(single_file_docs)
                        break  # Just process one file for testing
                
                if test_documents:
                    # Save test results
                    output_file = "./data/test_vector_db.json"
                    indexer.save_to_json(test_documents, output_file)
                    print(f"\n✅ Test completed! Processed {len(test_documents)} document chunks.")
                    print(f"📄 Test results saved to: {output_file}")
                else:
                    print("\n⚠️  No new files to process in test mode")
            
            elif choice == '2':
                # Full processing mode
                print("\n🚀 Full processing mode")
                folder_id = input("📁 Enter Google Drive folder ID (or press Enter for all files): ").strip()
                folder_id = folder_id if folder_id else None
                
                documents = indexer.process_files(folder_id)
                
                # Save to JSON
                output_file = "./data/vector_db.json"
                indexer.save_to_json(documents, output_file)
                
                print(f"\n✅ Full processing completed! Processed {len(documents)} document chunks.")
                print(f"📄 Vector database saved to: {output_file}")
            
            elif choice == '3':
                # Show cache info
                cache_info = indexer.get_cache_info()
                print(f"\n📊 Cache Information:")
                print(f"   Cached files: {cache_info['cached_files_count']}")
                print(f"   Cache file exists: {cache_info['cache_file_exists']}")
                print(f"   Cache location: {cache_info['cache_file_path']}")
                
                if cache_info['cached_files_count'] > 0:
                    print(f"\n📝 Recently cached files:")
                    for i, file_hash in enumerate(list(indexer.processed_files)[:5], 1):
                        print(f"   {i}. {file_hash}")
                    if len(indexer.processed_files) > 5:
                        print(f"   ... and {len(indexer.processed_files) - 5} more")
            
            elif choice == '4':
                # Clear cache
                confirm = input("\n⚠️  Are you sure you want to clear the cache? (y/N): ").strip().lower()
                if confirm == 'y':
                    indexer.clear_cache()
                    print("✅ Cache cleared successfully")
                else:
                    print("❌ Cache clear cancelled")
            
            elif choice == '5':
                print("\n👋 Goodbye!")
                break
            
            else:
                print("\n❌ Invalid choice. Please enter 1-5.")
    
    except Exception as e:
        logger.error(f"❌ Error: {e}")
        print(f"❌ An error occurred: {e}")

if __name__ == "__main__":
    main()
